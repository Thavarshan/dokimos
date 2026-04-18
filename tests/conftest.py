"""Shared test fixtures."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from dokimos.config import Settings
from dokimos.schemas.documents import Chunk, Document
from dokimos.schemas.index import (
    ChunkRef,
    IndexedChunk,
    IndexedSource,
    InvertedShingleIndex,
    SourceIndex,
)

# Text used inside binary fixture files so tests can assert on content.
FIXTURE_BODY = "The quick brown fox jumps over the lazy dog."


@pytest.fixture()
def sample_text() -> str:
    """A short sample document body for testing."""
    return (
        "The quick brown fox jumps over the lazy dog. "
        "Pack my box with five dozen liquor jugs. "
        "How vexingly quick daft zebras jump. "
        "The five boxing wizards jump quickly."
    )


@pytest.fixture()
def sample_file(tmp_path: Path, sample_text: str) -> Path:
    """Write ``sample_text`` to a temporary file and return its path."""
    p = tmp_path / "sample.txt"
    p.write_text(sample_text, encoding="utf-8")
    return p


@pytest.fixture()
def sample_document(sample_file: Path, sample_text: str) -> Document:
    """A pre-built :class:`Document` from the sample text."""
    return Document(
        source_path=sample_file,
        raw_text=sample_text,
        metadata={"filename": sample_file.name},
    )


@pytest.fixture()
def sample_chunks(sample_document: Document) -> list[Chunk]:
    """A minimal list of chunks derived from the sample document."""
    return [
        Chunk(
            document_id=sample_document.id,
            text=sample_document.raw_text,
            index=0,
            start_offset=0,
            end_offset=len(sample_document.raw_text),
        ),
    ]


@pytest.fixture()
def multi_paragraph_text() -> str:
    """Multi-paragraph sample for chunker tests."""
    return (
        "First paragraph about foxes and dogs.\n\n"
        "Second paragraph about liquor jugs and boxes.\n\n"
        "Third paragraph about zebras jumping quickly."
    )


@pytest.fixture()
def multi_sentence_text() -> str:
    """Multi-sentence sample for sentence-chunker tests."""
    return (
        "The quick brown fox jumps over the lazy dog. "
        "Pack my box with five dozen liquor jugs. "
        "How vexingly quick daft zebras jump."
    )


@pytest.fixture()
def corpus_dir(tmp_path: Path) -> Path:
    """A temporary corpus directory with a few .txt files."""
    d = tmp_path / "corpus"
    d.mkdir()
    (d / "source_a.txt").write_text("The fox jumped over the lazy dog.", encoding="utf-8")
    (d / "source_b.txt").write_text(
        "A completely different document about cats.", encoding="utf-8"
    )
    return d


@pytest.fixture()
def index_file(tmp_path: Path) -> Path:
    """Path for a temporary index file."""
    return tmp_path / "index.json"


@pytest.fixture()
def settings_with_index(tmp_path: Path, index_file: Path) -> Settings:
    """Settings configured to use tmp_path for corpus and index."""
    return Settings(
        corpus_path=tmp_path / "corpus",
        index_file=index_file,
        chunk_strategy="paragraph",
    )


@pytest.fixture()
def populated_index(index_file: Path) -> Path:
    """Write a small source index (v3 with inverted index) to disk."""
    shingles = [
        "the fox jumped over the",
        "fox jumped over the lazy",
        "jumped over the lazy dog.",
    ]
    idx = SourceIndex(
        version="3",
        sources=[
            IndexedSource(
                source_id="src-1",
                source_path="/tmp/source_a.txt",
                label="source_a.txt",
                chunks=[
                    IndexedChunk(
                        chunk_index=0,
                        text="The fox jumped over the lazy dog.",
                        shingles=shingles,
                    ),
                ],
            ),
        ],
        inverted=InvertedShingleIndex(
            entries={
                s: [ChunkRef(source_idx=0, chunk_idx=0)]
                for s in shingles
            }
        ),
    )
    index_file.write_text(
        json.dumps(idx.model_dump(mode="json"), indent=2),
        encoding="utf-8",
    )
    return index_file


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal .docx file containing ``FIXTURE_BODY``."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph(FIXTURE_BODY)
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal single-page PDF containing ``FIXTURE_BODY``."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)  # A4-ish
    page.insert_text((72, 72), FIXTURE_BODY)
    p = tmp_path / "sample.pdf"
    doc.save(str(p))
    doc.close()
    return p
