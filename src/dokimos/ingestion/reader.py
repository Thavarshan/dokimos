"""Document reading utilities with multi-format support."""

from __future__ import annotations

import logging
from pathlib import Path

from dokimos.exceptions import IngestionError
from dokimos.schemas.documents import Document

logger = logging.getLogger(__name__)

_READERS: dict[str, str] = {
    ".txt": "_read_text",
    ".md": "_read_text",
    ".docx": "_read_docx",
    ".pdf": "_read_pdf",
}


def read_document(path: Path) -> Document:
    """Read a file and return a :class:`Document`.

    Supports .txt, .md (plain-text), .docx (via python-docx), and .pdf (via pymupdf).

    Raises
    ------
    IngestionError
        If the file cannot be read or the format is unsupported.
    """
    resolved = path.resolve()
    if not resolved.is_file():
        raise IngestionError(f"Path is not a file: {resolved}")

    suffix = resolved.suffix.lower()
    reader_name = _READERS.get(suffix)
    if reader_name is None:
        raise IngestionError(
            f"Unsupported file format '{suffix}'. Supported: {', '.join(sorted(_READERS))}"
        )

    reader_fn = globals()[reader_name]
    raw_text: str = reader_fn(resolved)

    logger.info("Ingested document %s (%d characters)", resolved.name, len(raw_text))
    return Document(
        source_path=resolved,
        raw_text=raw_text,
        metadata={"filename": resolved.name, "format": suffix},
    )


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except (OSError, UnicodeDecodeError) as exc:
        raise IngestionError(f"Failed to read {path}: {exc}") from exc


def _read_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise IngestionError(
            "python-docx is required for .docx files. Install with: pip install dokimos-cli[docx]"
        ) from exc

    try:
        doc = DocxDocument(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        raise IngestionError(f"Failed to read DOCX {path}: {exc}") from exc


def _read_pdf(path: Path) -> str:
    try:
        import pymupdf
    except ImportError as exc:
        raise IngestionError(
            "pymupdf is required for .pdf files. Install with: pip install dokimos-cli[pdf]"
        ) from exc

    try:
        doc = pymupdf.open(str(path))
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(pages)
    except Exception as exc:
        raise IngestionError(f"Failed to read PDF {path}: {exc}") from exc
